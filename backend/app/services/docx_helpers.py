"""
Document generation helpers and utilities.

This module contains reusable functions for DOCX generation,
extracted from generator.py for better maintainability.
"""
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============ Font Helpers ============

def set_font(run, font_name='Times New Roman', font_size=12, bold=False, color=None):
    """Set font properties for a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


# ============ Table Helpers ============

def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_table_borders(table, style='single', size=4):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), style)
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_table_width(table, width_cm):
    """Set explicit table width in cm."""
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    # Convert cm to twips (1cm ≈ 567 twips)
    width_twips = int(width_cm * 567)
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def set_table_indent(table, indent=0):
    """Set table indentation."""
    tblPr = table._tbl.tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(indent))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)


def set_cell_padding(cell, padding_pt=1):
    """Set cell padding in points."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    # Convert points to twentieths (1pt = 20 twentieths)
    twips = str(int(padding_pt * 20))
    for side in ['top', 'bottom', 'left', 'right']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), twips)
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_row_cant_split(row):
    """Prevent row from splitting across pages."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None:
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)


# ============ Template Helpers ============

def get_template_path():
    """Get the template path, checking Docker and local locations."""
    paths = [
        "/app/app/templates/base_template.docx",
        os.path.join(os.path.dirname(__file__), "../templates/base_template.docx"),
    ]
    for path in paths:
        if os.path.exists(path):
            return path
    return None


def create_document_from_template():
    """
    Create a document from template or blank.
    
    Returns:
        tuple: (Document, using_template: bool)
    """
    template_path = get_template_path()
    if template_path:
        return Document(template_path), True
    return Document(), False


def setup_document_styles(document):
    """Configure default document styles."""
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)


def setup_page_margins(document, left=2.5, right=1.5, top=2.0, bottom=2.0):
    """Set page margins in cm."""
    section = document.sections[0]
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


# ============ Asset Path Helpers ============

def get_asset_path(filename):
    """Get path to an asset file (logo, stamp, signature)."""
    docker_path = f"/app/{filename}"
    if os.path.exists(docker_path):
        return docker_path
    
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    local_path = os.path.join(base_dir, filename)
    if os.path.exists(local_path):
        return local_path
    
    return None


def get_images_dir():
    """Get the images directory path."""
    if os.path.exists("/app/images"):
        return "/app/images"
    
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    return os.path.join(base_dir, "_изображения")


# ============ Part Type Helpers ============

def is_electronics_part(part):
    """
    Determine if a part is electronics/electromechanical.
    
    Criteria:
    - component_type == 'electronics'
    - material contains 'электро'
    - has specs dictionary
    """
    if hasattr(part, 'component_type') and part.component_type == 'electronics':
        return True
    if hasattr(part, 'material') and part.material and 'электро' in part.material.lower():
        return True
    if hasattr(part, 'specs') and part.specs:
        return True
    return False


def get_part_manufacturer(part, default_supplier, is_electronics=None):
    """
    Get manufacturer for a part.
    
    Electronics: use part.manufacturer if available
    Mechanical: use supplier from invoice
    """
    if is_electronics is None:
        is_electronics = is_electronics_part(part)
    
    if is_electronics and hasattr(part, 'manufacturer') and part.manufacturer:
        return part.manufacturer
    return default_supplier
