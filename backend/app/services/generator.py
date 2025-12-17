"""
DOCX Document Generator for Technical Descriptions and Letters.

Generates:
- Technical descriptions (Техническое описание)
- Non-insurance letters (Письмо о нестраховании)  
- Decision 130 notifications (Уведомление по Решению 130)
"""
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import os
from datetime import datetime

from app.db.models import Part
from app.services.parser import normalize_date
from app.services.s3 import s3_service
from app.services.docx_helpers import (
    set_font,
    remove_table_borders,
    add_table_borders,
    set_row_cant_split,
    get_asset_path,
    get_images_dir,
    is_electronics_part,
    get_part_manufacturer,
    create_document_from_template,
    setup_document_styles,
)


def create_header(document):
    section = document.sections[0]
    header = section.header
    
    # Clear existing content
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Create 3-column table for header
    # Width 16.5cm (Matches new printable width: 21 - 2.5 - 2.0 = 16.5)
    table = header.add_table(rows=2, cols=3, width=Cm(16.5))
    table.autofit = False
    remove_table_borders(table)
    
    # Columns: 40%, 30%, 30% of 16.5cm
    # 16.5 * 0.4 = 6.6
    # 16.5 * 0.3 = 4.95
    table.columns[0].width = Cm(6.6)
    table.columns[1].width = Cm(4.95)
    table.columns[2].width = Cm(4.95)
    
    # Set cell padding to 1pt (approx 0.035cm)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), '20') # 20 twentieths of a point = 1 point
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
    
    # Merge cells
    # Logo: Cell(0,0) + Cell(1,0) - Vertical merge
    cell_logo = table.cell(0, 0)
    cell_logo.merge(table.cell(1, 0))
    
    # Company Name: Cell(0,1) + Cell(0,2) - Horizontal merge
    cell_name = table.cell(0, 1)
    cell_name.merge(table.cell(0, 2))
    
    # Insert Logo
    logo_path = "/app/logo.png"
    if not os.path.exists(logo_path):
        logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "logo.png")

    if os.path.exists(logo_path):
        paragraph = cell_logo.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        # Width 5.5cm fits in 6.6cm column
        run.add_picture(logo_path, width=Cm(5.5)) 
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(5)
    
    # Company Name
    paragraph = cell_name.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('ООО "БСЛ-Лаб"')
    set_font(run, font_name='Calibri', font_size=14, bold=True)
    
    # Address (Row 2, Col 2)
    cell_addr = table.cell(1, 1)
    p_addr = cell_addr.paragraphs[0]
    p_addr.text = "ХМАО-Югра, г. Нижневартовск\nул. Кузоваткина, д. 3, ст. 4\nОГРН 1238600000143"
    p_addr.style.font.name = 'Calibri'
    p_addr.style.font.size = Pt(9)
    p_addr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_addr.paragraph_format.space_before = Pt(0)
    p_addr.paragraph_format.space_after = Pt(0)
    
    # Contacts (Row 2, Col 3)
    cell_contacts = table.cell(1, 2)
    p_contacts = cell_contacts.paragraphs[0]
    p_contacts.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_contacts.text = "ИНН 8603248341\nТел.: +7 9828801096\nE-mail: info@bsl-lab.ru"
    p_contacts.style.font.name = 'Calibri'
    p_contacts.style.font.size = Pt(9)
    p_contacts.paragraph_format.space_before = Pt(0)
    p_contacts.paragraph_format.space_after = Pt(0)

def create_footer(document):
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number
    run = paragraph.add_run()
    
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def create_outgoing_info(document, doc_number="80", doc_date=None, recipient="В таможенные органы"):
    if doc_date is None:
        doc_date = datetime.now().strftime("%d.%m.%Y")
    # REMOVED empty paragraphs as requested
    
    # Table for "Исх.№" and Recipient
    # Width 17.0cm
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    remove_table_borders(table)
    
    # 17.0cm total
    table.columns[0].width = Cm(11.0)
    table.columns[1].width = Cm(6.0)
    
    # Outgoing Info
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    run = p_left.add_run(f"Исх.№ {doc_number}   от   {doc_date}")
    run.font.underline = True
    set_font(run, font_size=12)
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Recipient
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.text = recipient
    p_right.style.font.name = 'Times New Roman'
    p_right.style.font.size = Pt(12)
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_float_picture(run, image_path, width, pos_x, pos_y):
    """
    Insert a floating image (In Front of Text) anchored to the paragraph.
    pos_x, pos_y: offsets in EMUs from the anchor point (paragraph).
    """
    # Add the picture normally first to get the relationship ID
    inline = run.add_picture(image_path, width=width)._inline
    
    # Get the graphic object
    graphic = inline.graphic
    
    # Create the anchor element
    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', "0")
    anchor.set('distB', "0")
    anchor.set('distL', "114300")
    anchor.set('distR', "114300")
    anchor.set('simplePos', "0")
    anchor.set('relativeHeight', "251658240")
    anchor.set('behindDoc', "0") # 0 = In Front of Text
    anchor.set('locked', "0")
    anchor.set('layoutInCell', "1")
    anchor.set('allowOverlap', "1")
    
    # Simple Positioning (coordinates)
    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', "0")
    simplePos.set('y', "0")
    anchor.append(simplePos)
    
    # Horizontal Position (Relative to Column or Page? Let's use Column/Margin)
    # To match the user's request "In Front of Text" and "Positioned", we use absolute positioning relative to Column
    positionH = OxmlElement('wp:positionH')
    positionH.set('relativeFrom', "column")
    posOffsetH = OxmlElement('wp:posOffset')
    posOffsetH.text = str(pos_x)
    positionH.append(posOffsetH)
    anchor.append(positionH)
    
    # Vertical Position (Relative to Paragraph)
    positionV = OxmlElement('wp:positionV')
    positionV.set('relativeFrom', "paragraph")
    posOffsetV = OxmlElement('wp:posOffset')
    posOffsetV.text = str(pos_y)
    positionV.append(posOffsetV)
    anchor.append(positionV)
    
    # Extent (Size)
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(inline.extent.cx))
    extent.set('cy', str(inline.extent.cy))
    anchor.append(extent)
    
    # Wrap None (In Front of Text behavior)
    wrapNone = OxmlElement('wp:wrapNone')
    anchor.append(wrapNone)
    
    # DocPr
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', '666') # Unique ID
    docPr.set('name', 'Floating Image')
    anchor.append(docPr)
    
    # Graphic
    anchor.append(graphic)
    
    # Replace inline with anchor
    inline.getparent().replace(inline, anchor)
    return anchor

def create_signature(document, add_facsimile=False):
    # VBA: 2 empty paragraphs before
    document.add_paragraph()
    document.add_paragraph()
    
    # VBA: Table 3 rows, 4 columns.
    table = document.add_table(rows=3, cols=4)
    table.autofit = False
    
    # Custom widths: 
    # Col 0 (Director): 4.0cm
    # Col 1 (Stamp): 8.0cm (Much Wider)
    # Col 2 (Name): 4.0cm
    # Col 3 (Rest): 1.0cm (Narrower)
    
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(8.0)
    table.columns[2].width = Cm(4.0)
    table.columns[3].width = Cm(1.0)
        
    # Set minimal padding
    
    # Row 0-1, Col 0: Director (Merged)
    cell_title = table.cell(0, 0)
    cell_title.merge(table.cell(1, 0))
    p_title = cell_title.paragraphs[0]
    p_title.text = "С уважением,\nДиректор"
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    cell_title.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_font(p_title.runs[0], font_size=12)
    
    # Row 0-1, Col 2: Name (Merged)
    cell_name = table.cell(0, 2) 
    cell_name.merge(table.cell(1, 2))
    p_name = cell_name.paragraphs[0]
    p_name.text = "А.В. Семенов"
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER # Changed to CENTER to match image
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(0)
    cell_name.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_font(p_name.runs[0], font_size=12)
    
    if add_facsimile:
        # Paths
        signature_path = "/app/signature.png"
        if not os.path.exists(signature_path):
            signature_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "signature.png")
            
        stamp_path = "/app/stamp.png"
        if not os.path.exists(stamp_path):
            stamp_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "stamp.png")

        # Check what we have
        has_stamp = os.path.exists(stamp_path)
        has_sig = os.path.exists(signature_path)
        
        # We will anchor images to the middle cell (Row 0, Col 1)
        # We merge Row 0 and 1 of Col 1 just to have a clean anchor point
        cell_anchor = table.cell(0, 1)
        cell_anchor.merge(table.cell(1, 1))
        p_anchor = cell_anchor.paragraphs[0]
        run = p_anchor.add_run()
        
        # Offsets in EMUs (1 cm = 360000 EMUs)
        # Col Width = 8.0cm = 2880000 EMUs
        
        if has_stamp:
            # Stamp: 100% scale (width=None)
            # Restore X: -324000 (Undo "Left 2cm")
            # Move DOWN: -360000 -> 0
            add_float_picture(run, stamp_path, width=None, pos_x=-324000, pos_y=0) 
            
        if has_sig:
            # Signature: 100% scale (width=None)
            # Restore Y: 0 (Undo "Down 1cm")
            # Move LEFT: 920000 -> 200000 (Shift left by ~2cm)
            add_float_picture(run, signature_path, width=None, pos_x=200000, pos_y=0)

def generate_technical_description(items: list, output_path: str, country_of_origin: str = "Китай", contract_no: str = None, contract_date: str = None, supplier: str = "Dongguan City Fangling Precision Mould Co., Ltd.", add_facsimile: bool = False):
    # Normalize date
    contract_date = normalize_date(contract_date)

    # Template logic
    TEMPLATE_PATH = "/app/app/templates/base_template.docx"
    if not os.path.exists(TEMPLATE_PATH):
        # Fallback to relative path for local testing if needed
        TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "../templates/base_template.docx")

    if os.path.exists(TEMPLATE_PATH):
        document = Document(TEMPLATE_PATH)
        using_template = True
    else:
        document = Document()
        using_template = False

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    if not using_template:
        create_header(document)
    
    create_footer(document)

    # Spacing from top header (2 lines)
    document.add_paragraph()
    document.add_paragraph()

    # Technical Description usually doesn't have the "Outgoing No" line in the same way, 
    # but user said "on the first page of EACH document". So I will add it.
    # But wait, Technical Description is an attachment usually. 
    # However, I will follow the instruction "row with date and number... on first page of each document".
    create_outgoing_info(document)

    # Spacing
    document.add_paragraph()
    
    # Title
    h1 = document.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run('Техническое описание.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = True
    
    # Subtitle / Description
    p_desc = document.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = (
        "Комплектующие предназначены для сборки Реометра Ротационного BSL R1, "
        "производства ООО «БСЛ-Лаб» (Россия), декларация о соответствии "
        "ЕАЭС N RU Д-RU.РА01.В. 89390/23, ТУ 26.51.53-001-82013305-2023)."
    )
    run = p_desc.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Contract Info (if needed, but usually this replaces it or goes below)
    # Based on image, it looks like this text IS the header/title block.
    # The original code had "Contract Info" below. I will keep it but maybe add spacing.
    # Contract Info removed as requested
    # if contract_no:
    #     contract_text = f"к Контракту № {contract_no}"
    #     if contract_date:
    #         contract_text += f" от {contract_date}"
    #     p_contract = document.add_paragraph(contract_text)
    #     p_contract.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     set_font(p_contract.runs[0], font_size=12)

    # Determine image directory
    if os.path.exists("/app/images"):
        IMAGES_DIR = "/app/images"
    else:
        BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
        IMAGES_DIR = os.path.join(BASE_DIR, "_изображения")
    
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def set_row_cant_split(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)

    for idx, part in enumerate(items, 1):
        p_head = document.add_paragraph(f"Позиция {idx}")
        set_font(p_head.runs[0], bold=True)
        p_head.paragraph_format.space_before = Pt(2)
        p_head.paragraph_format.space_after = Pt(2)
        p_head.paragraph_format.keep_with_next = True 
        
        table = document.add_table(rows=0, cols=2)
        # table.style = 'Table Grid' # Removed to avoid KeyError
        
        # Add borders manually
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') # 4 = 1/2 pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        table.autofit = False
        table.columns[0].width = Cm(6.0)
        table.columns[1].width = Cm(11.0)
            
        def add_row(label, value, keep_next=True):
            row = table.add_row()
            set_row_cant_split(row)
            
            cell_label = row.cells[0]
            cell_label.text = label
            cell_label.width = Cm(6.0)
            cell_label.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell_label.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if keep_next:
                    paragraph.paragraph_format.keep_with_next = True
                for run in paragraph.runs:
                    set_font(run, bold=True)
            
            cell_value = row.cells[1]
            cell_value.text = str(value) if value else ""
            cell_value.width = Cm(11.0)
            cell_value.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell_value.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if keep_next:
                    paragraph.paragraph_format.keep_with_next = True
                for run in paragraph.runs:
                    set_font(run, bold=False)

        # Определяем тип компонента (электроника, электромеханика и т.д.)
        is_electronics = (hasattr(part, 'component_type') and part.component_type == 'electronics') or \
                        (hasattr(part, 'material') and part.material and 'электро' in part.material.lower()) or \
                        (hasattr(part, 'specs') and part.specs)

        # Производитель: из БД только для электроники, иначе из инвойса (supplier)
        if is_electronics and hasattr(part, 'manufacturer') and part.manufacturer:
            item_manufacturer = part.manufacturer
        else:
            item_manufacturer = supplier

        add_row("Наименование", part.name)
        add_row("Обозначение", part.designation)
        add_row("Производитель", item_manufacturer)
        add_row("Страна происхождения", country_of_origin)

        if is_electronics:
            # Формат для электроники с разделом "Характеристики"
            add_row("Характеристики", "")  # Заголовок раздела
            
            if hasattr(part, 'specs') and part.specs:
                # Render specs from JSON (flexible format)
                for key, value in part.specs.items():
                    add_row(key, str(value))
            else:
                # Legacy format fallback
                if hasattr(part, 'current_type') and part.current_type:
                    add_row("Род тока", part.current_type)
                if hasattr(part, 'input_voltage') and part.input_voltage:
                    add_row("Входное напряжение, В", part.input_voltage)
                if hasattr(part, 'input_current') and part.input_current:
                    add_row("Входной ток, А", part.input_current)
                if hasattr(part, 'processor') and part.processor:
                    add_row("Процессор", part.processor)
                if hasattr(part, 'ram_kb') and part.ram_kb:
                    add_row("Оперативная память, Кб", str(part.ram_kb))
                if hasattr(part, 'rom_mb') and part.rom_mb:
                    add_row("Постоянная память, МБ", str(part.rom_mb))
            
            add_row("Материал изготовления", part.material)
            add_row("Размеры (мм)", part.dimensions)
            
            # Масса с единицей измерения
            if hasattr(part, 'weight') and part.weight:
                weight_unit = getattr(part, 'weight_unit', 'кг') or 'кг'
                add_row(f"Масса, {weight_unit}", str(part.weight))
        else:
            # Стандартный формат для механических деталей
            add_row("Размеры (мм)", part.dimensions)
            add_row("Материал изготовления", part.material)
            if hasattr(part, 'condition') and part.condition:
                add_row("Состояние", part.condition)
            if hasattr(part, 'description') and part.description:
                add_row("Технические характеристики", part.description)
            
        row_image = table.add_row()
        set_row_cant_split(row_image)
        
        cell_image = row_image.cells[0].merge(row_image.cells[1])
        cell_image.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        p_img_label = cell_image.paragraphs[0]
        p_img_label.text = f"Фото (эскиз) {idx}. {part.name or ''} {part.designation}"
        p_img_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img_label.paragraph_format.space_after = Pt(6)
        p_img_label.paragraph_format.keep_with_next = False 
        
        for run in p_img_label.runs:
            set_font(run, bold=True)
        
        if hasattr(part, 'image_path') and part.image_path:
            image_full_path = os.path.join(IMAGES_DIR, part.image_path)
            
            image_stream = None
            if os.path.exists(image_full_path):
                image_stream = open(image_full_path, 'rb')
            else:
                # Try S3
                s3_stream = s3_service.get_file(part.image_path)
                if s3_stream:
                    image_stream = s3_stream

            if image_stream:
                try:
                    import io
                    from PIL import Image
                    
                    # Read into memory to avoid issues with stream handling in PIL
                    img_data = image_stream.read()
                    if hasattr(image_stream, 'close'):
                        image_stream.close()

                    with Image.open(io.BytesIO(img_data)) as img:
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        
                        max_width_cm = 14.0
                        max_height_cm = 9.0
                        width, height = img.size
                        aspect_ratio = width / height
                        new_width = max_width_cm
                        new_height = new_width / aspect_ratio
                        
                        if new_height > max_height_cm:
                            new_height = max_height_cm
                            new_width = new_height * aspect_ratio

                        run_img = p_img_label.add_run()
                        run_img.add_break()
                        run_img.add_picture(img_byte_arr, width=Cm(new_width), height=Cm(new_height))
                        
                except Exception as e:
                    p_img_label.add_run(f"\n[Error loading image: {e}]")
            else:
                p_img_label.add_run("\n[Image file not found]")
        else:
            p_img_label.add_run("\n[No image available]")
        
        # Добавляем ТН ВЭД код для электроники (после фото)
        if is_electronics and hasattr(part, 'tnved_code') and part.tnved_code:
            tnved_text = f"\n\n{part.tnved_code}"
            if hasattr(part, 'tnved_description') and part.tnved_description:
                tnved_text += f" - {part.tnved_description}"
            run_tnved = p_img_label.add_run(tnved_text)
            set_font(run_tnved, bold=True, font_size=10)
            
    create_signature(document, add_facsimile)
            
    document.save(output_path)
    return output_path

def generate_non_insurance_letter(items: list, output_path: str, contract_no: str, contract_date: str, invoice_no: str, invoice_date: str, waybill_no: str = "", add_facsimile: bool = False):
    # Normalize dates
    contract_date = normalize_date(contract_date)
    invoice_date = normalize_date(invoice_date)

    # Template logic
    TEMPLATE_PATH = "/app/app/templates/base_template.docx"
    if not os.path.exists(TEMPLATE_PATH):
        # Fallback to relative path for local testing if needed
        TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "../templates/base_template.docx")

    if os.path.exists(TEMPLATE_PATH):
        document = Document(TEMPLATE_PATH)
        using_template = True
    else:
        document = Document()
        using_template = False

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman' 
    font.size = Pt(12)
    
    if not using_template:
        create_header(document)
    
    create_footer(document)
    
    # Spacing from top header (2 lines)
    document.add_paragraph()
    document.add_paragraph()
    
    create_outgoing_info(document)
    
    # Spacing
    document.add_paragraph()
    
    # Title
    p_title = document.add_paragraph('ПИСЬМО О НЕСТРАХОВАНИИ')
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p_title.runs[0], bold=True, font_size=14)
    
    # Body
    wb_num = waybill_no if waybill_no else "________________"
    text = f"Настоящим письмом сообщаем, что груз по накладной {wb_num} не страховался на всем пути следования."
    
    p_body = document.add_paragraph(text)
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body.paragraph_format.first_line_indent = Cm(1.25)
    
    create_signature(document, add_facsimile)
    document.save(output_path)
    return output_path

def generate_decision_130_notification(items: list, output_path: str, contract_no: str, contract_date: str, invoice_no: str, invoice_date: str, add_facsimile: bool = False):
    # Normalize dates
    contract_date = normalize_date(contract_date)
    invoice_date = normalize_date(invoice_date)

    # Template logic
    TEMPLATE_PATH = "/app/app/templates/base_template.docx"
    if not os.path.exists(TEMPLATE_PATH):
        # Fallback to relative path for local testing if needed
        TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "../templates/base_template.docx")

    if os.path.exists(TEMPLATE_PATH):
        document = Document(TEMPLATE_PATH)
        using_template = True
    else:
        document = Document()
        using_template = False
        
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman' 
    font.size = Pt(12)
    
    if not using_template:
        # Set margins (Left 2.5, Right 1.5 -> 17.0cm text width)
        section = document.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    if not using_template:
        create_header(document)
    
    create_footer(document)
    
    # If using template, header/footer are already there.
    # We might want to ensure footer is correct, but create_footer is empty anyway.
    
    # Spacing from top header (2 lines)
    document.add_paragraph()
    document.add_paragraph()
    
    create_outgoing_info(document)
    
    # Spacing (Only one empty line as requested)
    document.add_paragraph()
    
    # Title
    p_title = document.add_paragraph('УВЕДОМЛЕНИЕ')
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    set_font(p_title.runs[0], bold=True, font_size=14)
    
    # Subtitle
    subtitle_text = (
        "о подтверждении использования в заявленных нуждах и целях "
        "ввозимой (ввезенной) продукции, подлежащей обязательной "
        "оценке соответствия на таможенной территории Евразийского "
        "экономического союза, в отношении которой при помещении "
        "под таможенные процедуры не требуется подтверждение "
        "соблюдения мер технического регулирования"
    )
    p_subtitle = document.add_paragraph(subtitle_text)
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_subtitle.paragraph_format.first_line_indent = Cm(1.25)
    p_subtitle.paragraph_format.space_after = Pt(24)
    
    # Body Intro (Justified as requested "по ширине")
    intro_text = (
        "НАСТОЯЩИМ ООО «БСЛ-Лаб» ИНН 8603248341 УВЕДОМЛЯЕТ о ввозе на "
        "таможенную территорию Евразийского экономического союза продукции: "
        "части реометра ротационного BSL R1:"
    )
    p_intro = document.add_paragraph(intro_text)
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.paragraph_format.first_line_indent = Cm(1.25)
    p_intro.paragraph_format.space_after = Pt(12)
    
    # Table of Items
    table = document.add_table(rows=1, cols=3)
    # table.style = 'Table Grid' # Removed to avoid KeyError with custom template
    table.autofit = False
    
    # Force table alignment to Left
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Explicitly set table indentation to 0 to prevent shifting
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    tblPr = table._tbl.tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Add borders manually since we can't rely on 'Table Grid'
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 4 = 1/2 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # RADICAL FIX: Explicitly set table width to 17.0cm (approx 9638 dxa)
    # This ensures it matches the text width (21cm - 2.5cm - 1.5cm = 17cm)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9638') # 17.0cm in dxa (twips)
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Adjust column widths (Total 17.0cm)
    # Col 0: 5.0cm
    # Col 1: 9.0cm
    # Col 2: 3.0cm
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(9.0)
    table.columns[2].width = Cm(3.0)
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Обозначение'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[2].text = 'Кол-во'
    
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cell.paragraphs[0].runs[0], bold=True)
        # Add shading/background color
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9') # Light Gray
        tcPr.append(shd)
    
    # Data Rows
    for item in items:
        row_cells = table.add_row().cells
        
        # Designation
        # Handle both dictionary and object access for safety
        if isinstance(item, dict):
            designation = item.get('designation', '')
            name = item.get('name', '')
            quantity = item.get('quantity', '')
        else:
            designation = getattr(item, 'designation', '')
            name = getattr(item, 'name', '')
            quantity = getattr(item, 'quantity', '')

        row_cells[0].text = designation or ''
        p_des = row_cells[0].paragraphs[0]
        p_des.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_des.paragraph_format.left_indent = Cm(0.2)
        p_des.paragraph_format.space_before = Pt(1)
        p_des.paragraph_format.space_after = Pt(1)
        row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Name
        row_cells[1].text = name or ''
        p_name = row_cells[1].paragraphs[0]
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_name.paragraph_format.left_indent = Cm(0.2)
        p_name.paragraph_format.space_before = Pt(1)
        p_name.paragraph_format.space_after = Pt(1)
        row_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Quantity
        qty_str = "1"
        if quantity:
             qty_str = str(quantity)
        row_cells[2].text = qty_str
        p_qty = row_cells[2].paragraphs[0]
        p_qty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_qty.paragraph_format.space_before = Pt(1)
        p_qty.paragraph_format.space_after = Pt(1)
        row_cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
        row_cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    document.add_paragraph().paragraph_format.space_after = Pt(12)

    # Missing Text Part 1
    text_part_1 = (
        "подлежащей обязательной оценке соответствия на таможенной территории "
        "Евразийского экономического союза, в отношении которой не требуется "
        "подтверждение соблюдения мер технического регулирования"
    )
    p_1 = document.add_paragraph(text_part_1)
    p_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_1.paragraph_format.first_line_indent = Cm(1.25)
    p_1.paragraph_format.space_after = Pt(12)

    # Missing Text Part 2
    text_part_2 = (
        'СОГЛАСНО подпункту "е" пункта 6 Порядка ввоза на таможенную территорию '
        "Евразийского экономического союза продукции, подлежащей обязательной "
        "оценке соответствия на таможенной территории Евразийского экономического "
        "союза, утвержденного Решением Совета Евразийской экономической комиссии "
        "от 12 ноября 2021 г. N 130."
    )
    p_2 = document.add_paragraph(text_part_2)
    p_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_2.paragraph_format.first_line_indent = Cm(1.25)
    p_2.paragraph_format.space_after = Pt(24)
    
    create_signature(document, add_facsimile)
    
    document.save(output_path)
    return output_path
