import sys
import os
import re
from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

# Add backend to path
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

from app.db.models import Part
from app.core.config import settings

# Setup DB connection
engine = create_engine(settings.DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
db = SessionLocal()

docx_path = "_для бд/Техническое описание по 2023668805.docx"
images_dir = "_изображения"

def clean_filename(text):
    # Match the logic used in extraction
    return re.sub(r'[\\/*?:"<>|]', "", text).strip()

def update_images():
    print(f"Reading {docx_path}...")
    doc = Document(docx_path)
    
    # Map Designation -> Description from tables
    part_map = {}
    for table in doc.tables:
        data = {}
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                key = cells[0].text.strip()
                val = cells[1].text.strip()
                if key and val:
                    data[key] = val
        
        des = data.get('Обозначение')
        name = data.get('Наименование')
        if des and name:
            part_map[des] = name

    # Also need to map Description from "Фото (эскиз) X. Description" to Designation
    # Because the image filename is based on that description line, NOT the table Name.
    # In extraction script:
    # match = re.search(r"Фото \(эскиз\) (\d+)\.\s*(.*)", text)
    # desc = match.group(2).strip()
    # image_map[idx] = safe_desc
    
    # We need to link that `safe_desc` to `Designation`.
    # The DOCX structure is:
    # Text: "Фото (эскиз) X. Description"
    # Image
    # Text: "Позиция X" (or before)
    # Table with Designation
    
    # Let's try to correlate by order.
    # Or text content.
    # The description in "Фото..." usually contains the designation or part of it.
    
    print("Scanning document for image descriptions...")
    image_descriptions = []
    for p in doc.paragraphs:
        text = p.text.strip()
        match = re.search(r"Фото \(эскиз\) (\d+)\.\s*(.*)", text)
        if match:
            desc = match.group(2).strip()
            image_descriptions.append(desc)
            
    print(f"Found {len(image_descriptions)} image descriptions.")
    
    # We have 14 parts in tables. And 14 image descriptions.
    # Let's assume they are in the same order?
    # The tables are usually after the image? Or before?
    # In the dump:
    # 6: Позиция 1
    # 8: Фото ...
    # Table 1
    # So Order matches.
    
    parts_in_order = []
    for table in doc.tables:
        # Check if it's a part table
        is_part = False
        des = None
        for row in table.rows:
            if 'Обозначение' in row.cells[0].text:
                des = row.cells[1].text.strip()
                is_part = True
                break
        if is_part and des:
            parts_in_order.append(des)
            
    print(f"Found {len(parts_in_order)} parts in tables.")
    
    if len(parts_in_order) != len(image_descriptions):
        print("Warning: Count mismatch. Mapping might be off.")
    
    updated = 0
    for i, des in enumerate(parts_in_order):
        if i < len(image_descriptions):
            raw_desc = image_descriptions[i]
            safe_name = clean_filename(raw_desc)
            
            # Find the file with this name + extension
            found_file = None
            for ext in ['.png', '.jpeg', '.jpg']:
                fname = f"{safe_name}{ext}"
                fpath = os.path.join(images_dir, fname)
                if os.path.exists(fpath):
                    found_file = fname
                    break
            
            if found_file:
                print(f"Linking {des} -> {found_file}")
                part = db.query(Part).filter(Part.designation == des).first()
                if part:
                    part.image_path = found_file
                    updated += 1
                else:
                    print(f"  Part {des} not found in DB.")
            else:
                print(f"  Image file not found for {des} (Expected: {safe_name}.*)")
        else:
            print(f"  No image description for {des}")
            
    db.commit()
    print(f"Updated {updated} parts.")

if __name__ == "__main__":
    try:
        update_images()
    except Exception as e:
        print(f"Error: {e}")
