from docx import Document
import re

docx_path = "_для бд/Техническое описание по 2023668805.docx"

try:
    doc = Document(docx_path)
    print("--- DOCX TEXT DUMP ---")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"{i}: {text}")
            
    print("\n--- TABLE DUMP ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            print([cell.text.strip() for cell in row.cells])
            
except Exception as e:
    print(f"Error: {e}")
