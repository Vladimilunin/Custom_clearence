from docx import Document
import os

def create_template_from_docx(source_path, output_path):
    if not os.path.exists(source_path):
        print(f"Error: Source file not found at {source_path}")
        return

    document = Document(source_path)
    
    # 1. Clear Body Paragraphs
    # We iterate in reverse to avoid index shifting issues when deleting
    for i in range(len(document.paragraphs) - 1, -1, -1):
        p = document.paragraphs[i]
        p._element.getparent().remove(p._element)
        
    # 2. Clear Body Tables
    for i in range(len(document.tables) - 1, -1, -1):
        t = document.tables[i]
        t._element.getparent().remove(t._element)
        
    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    document.save(output_path)
    print(f"Template saved successfully to {output_path}")

if __name__ == "__main__":
    # Paths inside Docker container
    SOURCE_FILE = "/app/source_template.docx"
    OUTPUT_FILE = "/app/app/templates/base_template.docx"
    
    create_template_from_docx(SOURCE_FILE, OUTPUT_FILE)
