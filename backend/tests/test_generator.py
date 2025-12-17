"""
Tests for the generator module.
Tests document generation functionality.
"""
import pytest
import os
import tempfile
from docx import Document

from app.services.generator import (
    normalize_date,
    generate_technical_description,
    generate_non_insurance_letter,
    generate_decision_130_notification,
    set_font,
    remove_table_borders,
)


class SimplePart:
    """Simple part object for testing."""
    
    def __init__(self, **kwargs):
        self.designation = kwargs.get('designation', 'TEST-001')
        self.name = kwargs.get('name', 'Test Part')
        self.material = kwargs.get('material', 'Steel')
        self.weight = kwargs.get('weight', 1.5)
        self.weight_unit = kwargs.get('weight_unit', 'кг')
        self.dimensions = kwargs.get('dimensions', '100x50x25')
        self.description = kwargs.get('description', 'Test description')
        self.image_path = kwargs.get('image_path', None)
        self.manufacturer = kwargs.get('manufacturer', None)
        self.component_type = kwargs.get('component_type', None)
        self.specs = kwargs.get('specs', None)
        # Legacy fields
        self.current_type = kwargs.get('current_type', None)
        self.input_voltage = kwargs.get('input_voltage', None)
        self.input_current = kwargs.get('input_current', None)
        self.processor = kwargs.get('processor', None)
        self.ram_kb = kwargs.get('ram_kb', None)
        self.rom_mb = kwargs.get('rom_mb', None)
        self.tnved_code = kwargs.get('tnved_code', None)
        self.tnved_description = kwargs.get('tnved_description', None)
        self.condition = kwargs.get('condition', None)


class TestGenerateTechnicalDescription:
    """Tests for generate_technical_description function."""
    
    @pytest.fixture
    def output_path(self):
        """Create temporary file path for output."""
        fd, path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        yield path
        # Cleanup
        if os.path.exists(path):
            os.unlink(path)
    
    def test_generate_single_item(self, output_path):
        """Test generating document with single item."""
        parts = [SimplePart()]
        
        result = generate_technical_description(parts, output_path)
        
        assert result == output_path
        assert os.path.exists(output_path)
        
        # Verify document can be opened
        doc = Document(output_path)
        assert len(doc.tables) >= 1
    
    def test_generate_multiple_items(self, output_path):
        """Test generating document with multiple items."""
        parts = [
            SimplePart(designation='PART-001', name='Part One'),
            SimplePart(designation='PART-002', name='Part Two'),
            SimplePart(designation='PART-003', name='Part Three'),
        ]
        
        result = generate_technical_description(parts, output_path)
        
        assert os.path.exists(output_path)
        doc = Document(output_path)
        
        # Should have at least 3 position tables (one per item)
        assert len(doc.tables) >= 3
    
    def test_generate_with_contract_info(self, output_path):
        """Test generating with contract information."""
        parts = [SimplePart()]
        
        result = generate_technical_description(
            parts,
            output_path,
            country_of_origin="Китай",
            contract_no="123/2025",
            contract_date="01.01.2025",
            supplier="Test Supplier Co."
        )
        
        assert os.path.exists(output_path)
    
    def test_generate_with_electronics(self, output_path):
        """Test generating document with electronics component."""
        parts = [SimplePart(
            designation='ELEC-001',
            name='Microcontroller',
            material='Электроника',
            component_type='electronics',
            specs={
                'Род тока': 'Постоянный',
                'Напряжение': '5В',
                'Потребление': '100мА'
            },
            manufacturer='Texas Instruments',
            tnved_code='8542310001',
            tnved_description='Микросхемы электронные'
        )]
        
        result = generate_technical_description(parts, output_path)
        
        assert os.path.exists(output_path)
        doc = Document(output_path)
        
        # Document should contain the electronics specs
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        # Combined table cells text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += '\n' + cell.text
        
        assert 'Характеристики' in full_text or 'Род тока' in full_text
    
    def test_generate_empty_items_list(self, output_path):
        """Test generating with empty items list."""
        parts = []
        
        result = generate_technical_description(parts, output_path)
        
        # Should still create a valid document
        assert os.path.exists(output_path)
        doc = Document(output_path)
        assert doc is not None


class TestGenerateNonInsuranceLetter:
    """Tests for generate_non_insurance_letter function."""
    
    @pytest.fixture
    def output_path(self):
        """Create temporary file path for output."""
        fd, path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        yield path
        if os.path.exists(path):
            os.unlink(path)
    
    def test_generate_basic(self, output_path):
        """Test basic non-insurance letter generation."""
        parts = [SimplePart()]
        
        result = generate_non_insurance_letter(
            parts,
            output_path,
            contract_no="123/2025",
            contract_date="01.01.2025",
            invoice_no="INV-001",
            invoice_date="15.01.2025",
            waybill_no="WB-12345"
        )
        
        assert result == output_path
        assert os.path.exists(output_path)
        
        doc = Document(output_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        assert 'ПИСЬМО О НЕСТРАХОВАНИИ' in full_text or 'WB-12345' in full_text
    
    def test_generate_without_waybill(self, output_path):
        """Test generation without waybill number."""
        parts = [SimplePart()]
        
        result = generate_non_insurance_letter(
            parts,
            output_path,
            contract_no="123/2025",
            contract_date="01.01.2025",
            invoice_no="INV-001",
            invoice_date="15.01.2025"
        )
        
        assert os.path.exists(output_path)


class TestGenerateDecision130Notification:
    """Tests for generate_decision_130_notification function."""
    
    @pytest.fixture
    def output_path(self):
        """Create temporary file path for output."""
        fd, path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        yield path
        if os.path.exists(path):
            os.unlink(path)
    
    def test_generate_basic(self, output_path):
        """Test basic Decision 130 notification generation."""
        parts = [
            SimplePart(designation='R1.001', name='Деталь 1'),
            SimplePart(designation='R1.002', name='Деталь 2'),
        ]
        
        result = generate_decision_130_notification(
            parts,
            output_path,
            contract_no="123/2025",
            contract_date="01.01.2025",
            invoice_no="INV-001",
            invoice_date="15.01.2025"
        )
        
        assert result == output_path
        assert os.path.exists(output_path)
        
        doc = Document(output_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        assert 'УВЕДОМЛЕНИЕ' in full_text
    
    def test_items_table_present(self, output_path):
        """Test that items table is present in document."""
        parts = [
            SimplePart(designation='R1.001', name='Деталь 1'),
        ]
        
        result = generate_decision_130_notification(
            parts,
            output_path,
            contract_no="123/2025",
            contract_date="01.01.2025",
            invoice_no="INV-001",
            invoice_date="15.01.2025"
        )
        
        doc = Document(output_path)
        
        # Should have at least one table with items
        assert len(doc.tables) >= 1
        
        # Check that the items table has more than just a header row (header + data rows)
        # Find the table with the most rows (likely the items table)
        max_rows = max(len(table.rows) for table in doc.tables)
        assert max_rows >= 2  # At least header + 1 data row


class TestHelperFunctions:
    """Tests for helper functions in generator module."""
    
    def test_set_font(self):
        """Test set_font helper function."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Test text")
        
        set_font(run, font_name='Arial', font_size=14, bold=True)
        
        assert run.font.name == 'Arial'
        assert run.font.bold == True
    
    def test_remove_table_borders(self):
        """Test remove_table_borders helper function."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        
        # Should not raise an exception
        remove_table_borders(table)
        
        # Table should still exist
        assert len(table.rows) == 2
        assert len(table.columns) == 2
